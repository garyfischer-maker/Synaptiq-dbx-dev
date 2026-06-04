"""Generate the platform overview Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(0x8B, 0xA4, 0xBD)
AMBER = RGBColor(0xC8, 0x95, 0x6A)
DARK  = RGBColor(0x2D, 0x37, 0x48)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_HEX = "EEF3F8"
BLUE_HEX  = "8BA4BD"


def shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def hdr_row(table, headers):
    row = table.rows[0]
    for cell, h in zip(row.cells, headers):
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        shade(cell, BLUE_HEX)


def data_row(table, values, bold_first=False, alt=False):
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        cell.text = str(val)
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.bold = True
        if alt:
            shade(cell, LIGHT_HEX)


doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(2.0)

# ── TITLE ──────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("Synaptiq Data Profiling Tool")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BLUE

p2 = doc.add_paragraph()
r2 = p2.add_run("Product Overview  |  Executive Summary  |  Technical Reference")
r2.font.size = Pt(11); r2.font.color.rgb = AMBER

doc.add_paragraph("Version 1.0 POC  |  Azure Databricks  |  June 2026").runs[0].font.size = Pt(9)
doc.add_paragraph()

# ── SECTION 1: EXECUTIVE SUMMARY ───────────────────────────────────
h = doc.add_heading("What Is It?", level=1)
h.runs[0].font.color.rgb = BLUE

doc.add_paragraph(
    "The Synaptiq Data Profiling Tool is an intelligent data quality platform built on "
    "Azure Databricks. It automatically profiles healthcare data tables, detects anomalies "
    "and statistical drift between environments, and provides a conversational AI interface "
    "that lets users ask plain-English questions about their data quality — without writing SQL."
).runs[0].font.size = Pt(10)

doc.add_paragraph()
h2 = doc.add_heading("The Problem It Solves", level=1)
h2.runs[0].font.color.rgb = BLUE

doc.add_paragraph(
    "Healthcare data moves through multiple environments and changes constantly. "
    "Without automated quality checks, teams face hidden risks:"
).runs[0].font.size = Pt(10)

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
hdr_row(t, ["Problem", "Business Impact"])
problems = [
    ("Schema changes go undetected", "Broken pipelines, silent data loss"),
    ("Distribution shifts between environments", "Wrong analytics, incorrect clinical insights"),
    ("Manual data validation", "Hours of engineer time per release cycle"),
    ("No audit trail of data quality", "Compliance risk, no root-cause history"),
    ("Data quality hidden in notebooks", "Not accessible to analysts or business users"),
]
for i, row in enumerate(problems):
    data_row(t, row, bold_first=True, alt=i % 2 == 0)

doc.add_paragraph()
h3 = doc.add_heading("Key Capabilities", level=1)
h3.runs[0].font.color.rgb = BLUE

caps = [
    ("A/B Table Comparison",
     "Profile any two versions of a table side-by-side (TEST vs PROD, week vs week). "
     "Detects schema changes and statistical drift for every column."),
    ("Drift Detection",
     "PSI, KS test, Chi-square, Jensen-Shannon divergence. Verdict per column: "
     "stable / moderate / significant drift."),
    ("DQ Alerts",
     "Fires alerts for extreme nulls, constant columns, imbalanced distributions, "
     "high cardinality, skewed numeric columns — with severity ratings."),
    ("Persistent Governance Repository",
     "Every run writes to 5 Delta governance tables — a permanent, queryable record "
     "of data quality history."),
    ("AI/BI Dashboard",
     "Three-page Databricks dashboard: Schema Drift trends, Column Drift metrics, "
     "and single-run A vs B drill-down."),
    ("Ask Genie — Conversational AI",
     "Plain-English questions answered instantly: 'Which columns have significant drift?' "
     "Genie shows its SQL and returns the result."),
    ("Time-Series Load Filtering",
     "Filter any profile run to a specific nightly load batch, enabling "
     "28+ days of trend analysis."),
]
for title, desc in caps:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(title + ": ")
    r.bold = True; r.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

doc.add_paragraph()
h4 = doc.add_heading("Who Uses It", level=1)
h4.runs[0].font.color.rgb = BLUE

t2 = doc.add_table(rows=1, cols=2)
t2.style = "Table Grid"
hdr_row(t2, ["User", "How They Use It"])
users = [
    ("Data Engineers", "Run A/B comparisons before promoting code to production"),
    ("Data Quality Analysts", "Monitor alert trends and PSI drift signals across the data estate"),
    ("Clinical Informatics", "Ask Genie about data quality before trusting analytics outputs"),
    ("Data Governance / Compliance", "Audit trail of all quality runs stored in Delta governance tables"),
]
for i, row in enumerate(users):
    data_row(t2, row, bold_first=True, alt=i % 2 == 0)

doc.add_page_break()

# ── TOPOLOGY DIAGRAM ────────────────────────────────────────────────
h5 = doc.add_heading("Platform Topology", level=1)
h5.runs[0].font.color.rgb = BLUE

topo_text = (
    "USERS                    DATABRICKS APP                 UNITY CATALOG\n"
    "─────                    ──────────────                 ─────────────\n"
    "\n"
    "Data Engineers  ───▶    ┌──────────────────┐  ────▶   dev.prod_main_claims.*\n"
    "Data Analysts   ───▶    │  Synaptiq Data    │  ────▶   dev.test_main_claims.*\n"
    "Business Users  ───▶    │  Profiling Tool   │  ────▶   dev.prod_main_clinical.*\n"
    "                        │  (Streamlit App)  │  ────▶   dev.test_main_clinical.*\n"
    "                        └────────┬──────────┘  ────▶   dev.*_main_members.*\n"
    "                                 │\n"
    "              ┌──────────────────┼──────────────────┐\n"
    "              ▼                  ▼                  ▼\n"
    "    ┌─────────────────┐  ┌──────────────┐  ┌──────────────────────┐\n"
    "    │  SQL Warehouse  │  │  UC Volume   │  │  Governance Delta    │\n"
    "    │  Statement      │  │  ab_runs/    │  │  Tables              │\n"
    "    │  Execution API  │  │  HTML/Excel  │  │  profiler_runs       │\n"
    "    └─────────────────┘  │  Mermaid/JSON│  │  dataset_profiles    │\n"
    "                          └──────────────┘  │  column_profiles     │\n"
    "                                            │  column_alerts       │\n"
    "                                            │  column_comparisons  │\n"
    "                                            └──────────┬───────────┘\n"
    "                                                       │\n"
    "                                         ┌─────────────┴──────────────┐\n"
    "                                         ▼                            ▼\n"
    "                                ┌──────────────────┐    ┌────────────────────┐\n"
    "                                │  AI/BI Dashboard │    │  Genie AI Space    │\n"
    "                                │  Schema Drift    │    │  Conversational Q&A│\n"
    "                                │  Column Drift    │    │  over DQ data      │\n"
    "                                │  Run Comparison  │    │  (Public Preview)  │\n"
    "                                └──────────────────┘    └────────────────────┘\n"
)
p_topo = doc.add_paragraph(topo_text)
p_topo.runs[0].font.name = "Courier New"
p_topo.runs[0].font.size = Pt(7)

doc.add_paragraph()
h6 = doc.add_heading("Key Outputs Per Run", level=2)
h6.runs[0].font.color.rgb = BLUE

t3 = doc.add_table(rows=1, cols=3)
t3.style = "Table Grid"
hdr_row(t3, ["Output", "Format", "Location"])
outs = [
    ("HTML Profile Reports", "Interactive web page (per side)", "UC Volume"),
    ("Excel Workbook", "5 sheets: Schema Diff, Metrics, Alerts, Drift, Row Diff", "UC Volume"),
    ("Mermaid Diagrams", "Side A schema, Side B schema, Drift-only view", "UC Volume"),
    ("Metamodel JSON", "Machine-readable DQ record", "UC Volume"),
    ("Governance Records", "5 Delta tables — full queryable history", "Unity Catalog"),
]
for i, row in enumerate(outs):
    data_row(t3, row, alt=i % 2 == 0)

doc.add_page_break()

# ── SECTION 2: TECHNICAL REFERENCE ─────────────────────────────────
p_tech = doc.add_paragraph()
r_tech = p_tech.add_run("Technical Reference")
r_tech.bold = True; r_tech.font.size = Pt(18); r_tech.font.color.rgb = BLUE

h7 = doc.add_heading("Technology Stack", level=1)
h7.runs[0].font.color.rgb = BLUE

t4 = doc.add_table(rows=1, cols=2)
t4.style = "Table Grid"
hdr_row(t4, ["Layer", "Technology"])
stack = [
    ("UI Framework", "Streamlit (Python 3.11)"),
    ("Hosting", "Databricks Apps — containerised serverless compute"),
    ("Catalog & Security", "Unity Catalog — RLS/CLS enforced at UC layer"),
    ("Data Storage", "UC Volumes (artifacts) + Delta Lake (governance tables)"),
    ("SQL Execution", "Databricks Statement Execution API — no JDBC"),
    ("Metadata Lookups", "Databricks SDK WorkspaceClient — UC REST API"),
    ("Drift Metrics", "SciPy (KS, Chi-square) + NumPy (PSI, JS divergence)"),
    ("Conversational AI", "Databricks Genie Conversation API (Public Preview)"),
    ("Identity", "OAuth M2M — app SP; UC enforces per-table grants"),
]
for i, row in enumerate(stack):
    data_row(t4, row, bold_first=True, alt=i % 2 == 0)

doc.add_paragraph()
h8 = doc.add_heading("Governance Delta Schema  (dev.test_main_profiler)", level=1)
h8.runs[0].font.color.rgb = BLUE

t5 = doc.add_table(rows=1, cols=3)
t5.style = "Table Grid"
hdr_row(t5, ["Table", "Grain", "Key Columns"])
schema = [
    ("profiler_runs", "1 row / run",
     "run_id (PK), run_label, created_utc, side_a_fqn, side_b_fqn, lineage_json"),
    ("dataset_profiles", "2 rows / run",
     "run_id, side (A|B), env_label, catalog, schema, table, row_count, column_count"),
    ("column_profiles", "1 row / column / side / run",
     "run_id, side, column_name, logical_type, null_pct, distinct_pct, numeric stats, cat stats, stereotypes"),
    ("column_alerts", "1 row / alert",
     "run_id, side, column_name, rule, severity (critical|warn|info), message"),
    ("column_comparisons", "1 row / column / run",
     "run_id, column_name, schema_change, PSI, KS stat, chi_square, js_divergence, verdict, stereotypes"),
]
for i, row in enumerate(schema):
    data_row(t5, row, bold_first=True, alt=i % 2 == 0)

doc.add_paragraph()
h9 = doc.add_heading("Key Design Decisions", level=1)
h9.runs[0].font.color.rgb = BLUE

decisions = [
    ("Statement Execution API (not JDBC)",
     "JDBC sql.connect() hangs indefinitely in Apps containers. REST API uses same "
     "OAuth path as catalog lookups — already confirmed working."),
    ("UC Files API (not FUSE mount)",
     "FUSE mount (/Volumes/...) unavailable in Apps container. Files API writes "
     "directly to UC Volumes via REST — no mount required."),
    ("Custom pandas profiling (not ydata-profiling)",
     "ydata-profiling adds 2-5 min startup even on tiny tables. Custom pandas "
     "stats complete in seconds for the same column metrics."),
    ("Pure Python drift from histograms",
     "PSI, KS, Chi-sq, and JS divergence computed from pre-computed histograms — "
     "no second warehouse query needed per run."),
    ("UC REST for schema discovery",
     "Warehouse cold-start (2-5 min) blocked the UI dropdowns. UC REST returns "
     "schema/table lists in <1 second with no running warehouse."),
]
for title, desc in decisions:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(title + ": ")
    r.bold = True; r.font.size = Pt(9)
    p.add_run(desc).font.size = Pt(9)

doc.add_paragraph()
h10 = doc.add_heading("Deployment Reference", level=1)
h10.runs[0].font.color.rgb = BLUE

t6 = doc.add_table(rows=1, cols=2)
t6.style = "Table Grid"
hdr_row(t6, ["Component", "Value"])
deploy = [
    ("App URL", "https://synaptiq-dq-platform-7405619521761591.11.azuredatabricks.net/"),
    ("SQL Warehouse", "2ad65b4df5cd3a9e (serverless)"),
    ("Output Volume", "dev.test_main_profiler.ab_runs"),
    ("App Service Principal", "39ee93a7-c623-4614-90a8-c3798bb5b329"),
    ("Genie Space ID", "01f15f8e849e1c31a47c5785d07751bf"),
    ("AI/BI Dashboard", "/dashboardsv3/01f15f7d18171dac85cdc247e47d48a5"),
    ("Source Repository", "github.com/garyfischer-maker/Synaptiq-dbx-dev (main)"),
]
for i, row in enumerate(deploy):
    data_row(t6, row, bold_first=True, alt=i % 2 == 0)

# Footer
doc.add_paragraph()
p_foot = doc.add_paragraph()
p_foot.add_run("© 2026 Synaptiq  |  Confidential  |  synaptiq.ai").font.size = Pt(8)
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save("docs/synaptiq_dq_platform_overview.docx")
print("Saved: docs/synaptiq_dq_platform_overview.docx")
